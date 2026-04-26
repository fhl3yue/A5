import json
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from sqlalchemy import delete
from sqlalchemy.orm import Session

from app.models import KnowledgeChunk, KnowledgeDocument, RoutePreset, ScenicSpot
from app.utils import normalize_text


def load_json(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def upsert_sample_data(db: Session, sample_dir: Path) -> None:
    spots_path = sample_dir / "scenic_spots.json"
    routes_path = sample_dir / "routes.json"
    if not spots_path.exists() or not routes_path.exists():
        raise FileNotFoundError("Missing sample scenic data files.")

    db.execute(delete(ScenicSpot))
    db.execute(delete(RoutePreset))
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_name == "sample_scenic_spots"))
    db.execute(delete(KnowledgeDocument).where(KnowledgeDocument.name == "sample_scenic_spots"))

    spots = load_json(spots_path)
    for spot in spots:
        db.add(
            ScenicSpot(
                spot_id=spot["spot_id"],
                name=spot["name"],
                location=spot.get("location", ""),
                cultural_meaning=spot.get("cultural_meaning", ""),
                description=spot.get("description", ""),
                highlights=spot.get("highlights", ""),
                schedule=spot.get("schedule", ""),
                tags=",".join(spot.get("tags", [])),
            )
        )
        chunk_text = "；".join(
            [
                f"景点名称：{spot['name']}",
                f"位置：{spot.get('location', '')}",
                f"文化内涵：{spot.get('cultural_meaning', '')}",
                f"详细介绍：{spot.get('description', '')}",
                f"游玩亮点：{spot.get('highlights', '')}",
                f"开放或演出信息：{spot.get('schedule', '')}",
            ]
        )
        db.add(
            KnowledgeChunk(
                document_name="sample_scenic_spots",
                title=spot["name"],
                content=normalize_text(chunk_text),
                tags=",".join(spot.get("tags", [])),
            )
        )

    db.add(
        KnowledgeDocument(
            name="sample_scenic_spots",
            source="sample",
            status="active",
            content_type="json",
        )
    )

    routes = load_json(routes_path)
    for route in routes:
        db.add(
            RoutePreset(
                name=route["name"],
                interest=route["interest"],
                duration=route["duration"],
                spots="|".join(route["spots"]),
                reason=route["reason"],
            )
        )

    db.commit()


def import_plain_text_document(db: Session, file_path: Path, source: str = "upload") -> int:
    text = file_path.read_text(encoding="utf-8", errors="ignore")
    chunks = [normalize_text(item) for item in text.splitlines() if normalize_text(item)]
    document_name = file_path.name
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_name == document_name))
    db.execute(delete(KnowledgeDocument).where(KnowledgeDocument.name == document_name))

    db.add(KnowledgeDocument(name=document_name, source=source, status="active", content_type=file_path.suffix.lower()))
    for idx, chunk in enumerate(chunks, start=1):
        title = chunk[:24]
        db.add(
            KnowledgeChunk(
                document_name=document_name,
                title=title,
                content=chunk,
                tags="imported",
            )
        )
    db.commit()
    return len(chunks)


def import_docx_document(db: Session, file_path: Path, source: str = "official-docx") -> int:
    doc = Document(file_path)
    chunks = [normalize_text(p.text) for p in doc.paragraphs if normalize_text(p.text)]
    document_name = file_path.name
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_name == document_name))
    db.execute(delete(KnowledgeDocument).where(KnowledgeDocument.name == document_name))
    db.add(KnowledgeDocument(name=document_name, source=source, status="active", content_type="docx"))
    for idx, chunk in enumerate(chunks, start=1):
        title = chunk[:24]
        db.add(
            KnowledgeChunk(
                document_name=document_name,
                title=title,
                content=chunk,
                tags="docx",
            )
        )
    db.commit()
    return len(chunks)


def import_xlsx_rows(db: Session, file_path: Path, source: str = "official-xlsx", row_limit: int = 300) -> int:
    wb = load_workbook(file_path, read_only=True, data_only=True)
    ws = wb[wb.sheetnames[0]]
    rows = ws.iter_rows(values_only=True)
    headers = [str(item).strip() if item is not None else "" for item in next(rows)]
    imported = 0
    db.execute(delete(KnowledgeChunk).where(KnowledgeChunk.document_name == file_path.name))
    db.execute(delete(KnowledgeDocument).where(KnowledgeDocument.name == file_path.name))
    db.add(KnowledgeDocument(name=file_path.name, source=source, status="active", content_type="xlsx"))
    for idx, row in enumerate(rows, start=1):
        if imported >= row_limit:
            break
        row_map = {headers[i]: row[i] for i in range(min(len(headers), len(row)))}
        attraction_name = str(row_map.get("attraction_name", "") or "").strip()
        attraction_content = str(row_map.get("attraction_content", "") or "").strip()
        attraction_type = str(row_map.get("attraction_type", "") or "").strip()
        if not attraction_name or not attraction_content:
            continue
        text = normalize_text(
            f"景点名称：{attraction_name}；景点类型：{attraction_type}；介绍内容：{attraction_content[:1200]}"
        )
        db.add(
            KnowledgeChunk(
                document_name=file_path.name,
                title=attraction_name or f"{file_path.name}-row-{idx}",
                content=text,
                tags=f"xlsx,{attraction_type}",
            )
        )
        imported += 1
    db.commit()
    return imported
